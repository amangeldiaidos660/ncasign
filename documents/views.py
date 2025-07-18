from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.conf import settings
from django.contrib.auth import get_user_model
import os
from docx import Document as DocxDocument
from .forms import GPHContractForm, ActForm
from .models import GphDocument, ActDocument, ActPackage
from datetime import datetime
import dropbox
import io
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from django.views.decorators.http import require_GET
import json
from ncasign.utils import get_dropbox_access_token
from docx.table import _Row

# Универсальный декоратор для проверки ролей (копия из ncasign/views.py)
from functools import wraps

def role_required(roles):
    def decorator(view_func):
        @wraps(view_func)
        def _wrapped_view(request, *args, **kwargs):
            if request.user.role not in roles:
                return render(request, '403.html', status=403)
            return view_func(request, *args, **kwargs)
        return _wrapped_view
    return decorator

@login_required
@role_required([1, 5, 2])  # Админ, Редактор, Подписант (только просмотр)
def acts_list(request):
    return render(request, 'documents/acts_list.html')

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут создавать документы
def gph_create(request):
    """Создание нового ГПХ договора"""
    if request.method == 'POST':
        form = GPHContractForm(request.POST)
        if form.is_valid():
            # Здесь будет логика обработки формы
            pass
    else:
        form = GPHContractForm()
    
    return render(request, 'documents/gph_create.html', {
        'form': form
    })

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут создавать акты
def act_create(request):
    """Создание акта с формой и предпросмотром шаблона"""
    if request.method == 'GET':
        request.session['package_acts'] = []
    if request.method == 'POST':
        form = ActForm(request.POST)
        if form.is_valid():
            # Здесь будет логика обработки формы
            pass
    else:
        form = ActForm()
    
    # Получаем данные пользователей для JavaScript
    User = get_user_model()
    users_data = {}
    for user in User.objects.filter(role=4):
        users_data[user.username] = {
            'full_name': user.full_name or '',
            'phone_number': user.phone_number or '',
            'iin': user.iin or '',
        }
    
    import json
    users_data_json = json.dumps(users_data, ensure_ascii=False)
    
    return render(request, 'documents/act_create.html', {
        'form': form,
        'users_data_json': users_data_json
    })

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут делать предпросмотр
def gph_preview(request):
    if request.method == 'POST':
        data = {
            'full_name': '',
            'start_date': request.POST.get('start_date', ''),
            'end_date': request.POST.get('end_date', ''),
        }
        executor_id = request.POST.get('executor', '')
        if executor_id:
            User = get_user_model()
            try:
                executor = User.objects.get(username=executor_id, role=4)
                data['full_name'] = executor.full_name
            except User.DoesNotExist:
                pass
        signer_id = request.POST.get('signer', '')
        signer_full_name = ''
        proxy_number = ''
        proxy_date = ''
        if signer_id:
            User = get_user_model()
            try:
                signer = User.objects.get(username=signer_id, role=2)
                signer_full_name = signer.full_name
                proxy_number = signer.proxy_number or ''
                proxy_date = signer.proxy_date.strftime('%d.%m.%Y') if signer.proxy_date else ''
            except User.DoesNotExist:
                pass
        try:
            template_path = os.path.join(settings.BASE_DIR, 'static', 'docs', 'gph.docx')
            if os.path.exists(template_path):
                doc = DocxDocument(template_path)
                def replace_placeholders(text):
                    text = text.replace('{{full_name}}', f'<strong>{data["full_name"]}</strong>' if data['full_name'] else '')
                    text = text.replace('{{doc_id}}', '<em>будет сгенерирован при сохранении</em>')
                    text = text.replace('{{year}}', '<strong>2025</strong>')
                    text = text.replace('{{current_date}}', '<strong>13.07.2025</strong>')
                    text = text.replace('{{start_date}}', f'<strong>{data["start_date"]}</strong>' if data['start_date'] else '')
                    text = text.replace('{{end_date}}', f'<strong>{data['end_date']}</strong>' if data['end_date'] else '')
                    text = text.replace('{{nca_datas}}', '<strong>Данные подписи НУЦ</strong>')
                    text = text.replace('{{signer}}', f'<strong>{signer_full_name}</strong>' if signer_full_name else '')
                    text = text.replace('{{signer_num}}', f'<strong>{proxy_number}</strong>' if proxy_number else '')
                    text = text.replace('{{signer_date}}', f'<strong>{proxy_date}</strong>' if proxy_date else '')
                    return text
                for paragraph in doc.paragraphs:
                    paragraph.text = replace_placeholders(paragraph.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.text = replace_placeholders(paragraph.text)
                preview_html = '<div style="padding: 20px; background: white; border-radius: 10px; font-family: Arial, sans-serif;">'
                preview_html += '<div style="border: 1px solid #ddd; padding: 20px; background: #f9f9f9;">'
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        processed_text = replace_placeholders(paragraph.text)
                        preview_html += f'<p style="margin-bottom: 10px;">{processed_text}</p>'
                for table in doc.tables:
                    preview_html += '<table style="width:100%; margin-bottom: 20px; border-collapse: collapse;">'
                    for row in table.rows:
                        preview_html += '<tr>'
                        for cell in row.cells:
                            cell_texts = []
                            for p in cell.paragraphs:
                                if p.text.strip():
                                    processed_text = replace_placeholders(p.text)
                                    cell_texts.append(processed_text)
                            cell_text = '<br>'.join(cell_texts)
                            preview_html += f'<td style="border:1px solid #ccc; padding:5px;">{cell_text}</td>'
                        preview_html += '</tr>'
                    preview_html += '</table>'
                preview_html += '</div>'
                preview_html += '</div>'
                return JsonResponse({'success': True, 'preview_html': preview_html})
            else:
                preview_html = f"""
                <div style=\"padding: 20px; background: white; border-radius: 10px;\">
                    <h3>Файл шаблона gph.docx не найден</h3>
                </div>
                """
                return JsonResponse({'success': True, 'preview_html': preview_html})
        except Exception as e:
            return JsonResponse({'success': False, 'error': f'Ошибка загрузки документа: {str(e)}'})
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут делать предпросмотр акта
def act_preview(request):
    """Предпросмотр шаблона акта с подстановкой данных"""
    try:
        html_path = os.path.join(settings.BASE_DIR, 'static', 'docs', 'act.html')
        if not os.path.exists(html_path):
            return JsonResponse({'success': False, 'error': 'Файл шаблона act.html не найден'})
        with open(html_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        if request.method == 'POST':
            data = {
                'full_name': request.POST.get('full_name', ''),
                'phone_number': request.POST.get('phone_number', ''),
                'iin': request.POST.get('iin', ''),
                'doc_id': request.POST.get('doc_id', ''),
                'act_id': '',
                'current_date': timezone.now().strftime('%d.%m.%Y'),
                'start_date': request.POST.get('start_date', ''),
                'end_date': request.POST.get('end_date', ''),
                'quantity': request.POST.get('quantity', ''),
                'unit_price': request.POST.get('unit_price', ''),
                'amount': request.POST.get('amount', ''),
                'created_date': '',
                'text': request.POST.get('text', ''),
                'unit': request.POST.get('unit', ''),
                'additional_text': request.POST.get('additional_text', ''),
                'vat_included': request.POST.get('vat_included', '') == 'on',
            }
            # Форматируем даты для предпросмотра
            from datetime import datetime
            def format_date(val):
                try:
                    if isinstance(val, str) and val:
                        return datetime.strptime(val, '%Y-%m-%d').strftime('%d.%m.%Y') if '-' in val else val
                    return val
                except Exception:
                    return val
            data['start_date'] = format_date(data['start_date'])
            data['end_date'] = format_date(data['end_date'])
            # Если указан исполнитель, ищем его последний ГПХ договор только для ID
            executor_username = request.POST.get('executor', '')
            if executor_username and not data['doc_id']:
                try:
                    User = get_user_model()
                    user = User.objects.get(username=executor_username, role=4)
                    last_gph = GphDocument.objects.filter(
                        user=user, 
                        doc_type='gph'
                    ).order_by('-created_at').first()
                    if last_gph:
                        data['doc_id'] = last_gph.doc_id
                        data['created_date'] = last_gph.created_at.strftime('%d.%m.%Y') if last_gph.created_at else ''
                except User.DoesNotExist:
                    pass
            # Обрабатываем массив работ
            import json
            works = []
            works_json = request.POST.get('works', '[]')
            print(f"Получены данные работ: {works_json}")
            try:
                works = json.loads(works_json)
                print(f"Работы после парсинга: {works}")
            except json.JSONDecodeError as e:
                print(f"Ошибка парсинга JSON: {e}")
                works = []
            
            # Если работ нет, используем данные из формы как одну работу
            if not works:
                try:
                    quantity = float(data['quantity']) if data['quantity'] else 0
                    unit_price = float(data['unit_price']) if data['unit_price'] else 0
                    amount = quantity * unit_price
                    if data['vat_included']:
                        amount_vat = amount * 1.12
                        data['amount'] = f"{amount:.2f} (с НДС: {amount_vat:.2f})"
                    else:
                        data['amount'] = f"{amount:.2f}"
                except (ValueError, ZeroDivisionError):
                    data['amount'] = ''
                
                # Подставляем данные в HTML для одной работы
                html_content = html_content.replace('{{full_name}}', f'<strong>{data["full_name"]}</strong>' if data['full_name'] else '')
                html_content = html_content.replace('{{phone_number}}', f'<strong>{data["phone_number"]}</strong>' if data['phone_number'] else '')
                html_content = html_content.replace('{{iin}}', f'<strong>{data["iin"]}</strong>' if data['iin'] else '')
                html_content = html_content.replace('{{doc_id}}', f'<strong>{data["doc_id"]}</strong>' if data['doc_id'] else '')
                html_content = html_content.replace('{{created_date}}', f'<strong>{data["created_date"]}</strong>' if data['created_date'] else '')
                html_content = html_content.replace('{{act_id}}', f'<strong>{data["act_id"]}</strong>' if data['act_id'] else '')
                html_content = html_content.replace('{{current_date}}', f'<strong>{data["current_date"]}</strong>')
                html_content = html_content.replace('{{start_date}}', f'<strong>{data["start_date"]}</strong>' if data['start_date'] else '')
                html_content = html_content.replace('{{end_date}}', f'<strong>{data["end_date"]}</strong>' if data['end_date'] else '')
                html_content = html_content.replace('{{quantity}}', f'<strong>{data["quantity"]}</strong>' if data['quantity'] else '')
                html_content = html_content.replace('{{sum}}', f'<strong>{data["unit_price"]}</strong>' if data['unit_price'] else '')
                html_content = html_content.replace('{{amount}}', f'<strong>{data["amount"]}</strong>' if data['amount'] else '')
                html_content = html_content.replace('{{text}}', f'<strong>{data["text"]}</strong>' if data['text'] else '')
                html_content = html_content.replace('{{unit}}', f'<strong>{data["unit"]}</strong>' if data['unit'] else '')
                html_content = html_content.replace('{{additional_text}}', f'<strong>{data["additional_text"]}</strong>' if data['additional_text'] else '')
                # Итоговые значения
                res_quantity = data['quantity'] if data['quantity'] else ''
                res_sum = data['unit_price'] if data['unit_price'] else ''
                res_amount = data['amount'] if data['amount'] else ''
                
                # Показываем как целые числа, если возможно
                try:
                    if res_quantity and float(res_quantity).is_integer():
                        res_quantity = str(int(float(res_quantity)))
                    if res_sum and float(res_sum).is_integer():
                        res_sum = str(int(float(res_sum)))
                except (ValueError, TypeError):
                    pass
                
                html_content = html_content.replace('{{res_quantity}}', f'<strong>{res_quantity}</strong>')
                html_content = html_content.replace('{{res_sum}}', f'<strong>{res_sum}</strong>')
                html_content = html_content.replace('{{res}}', f'<strong>{res_amount}</strong>')
            else:
                # Подставляем основные данные
                html_content = html_content.replace('{{full_name}}', f'<strong>{data["full_name"]}</strong>' if data['full_name'] else '')
                html_content = html_content.replace('{{phone_number}}', f'<strong>{data["phone_number"]}</strong>' if data['phone_number'] else '')
                html_content = html_content.replace('{{iin}}', f'<strong>{data["iin"]}</strong>' if data['iin'] else '')
                html_content = html_content.replace('{{doc_id}}', f'<strong>{data["doc_id"]}</strong>' if data['doc_id'] else '')
                html_content = html_content.replace('{{created_date}}', f'<strong>{data["created_date"]}</strong>' if data['created_date'] else '')
                html_content = html_content.replace('{{act_id}}', f'<strong>{data["act_id"]}</strong>' if data['act_id'] else '')
                html_content = html_content.replace('{{current_date}}', f'<strong>{data["current_date"]}</strong>')
                
                # Создаем HTML для всех работ
                works_html = ''
                total_quantity = 0
                total_sum = 0
                total_amount = 0.0
                
                for i, work in enumerate(works):
                    # Форматируем даты
                    start_date = work.get("start_date", "")
                    end_date = work.get("end_date", "")
                    if start_date and end_date:
                        try:
                            from datetime import datetime
                            start_date = datetime.strptime(start_date, '%Y-%m-%d').strftime('%d.%m.%Y') if '-' in start_date else start_date
                            end_date = datetime.strptime(end_date, '%Y-%m-%d').strftime('%d.%m.%Y') if '-' in end_date else end_date
                        except:
                            pass
                    
                    works_html += f'''
                    <tr>
                        <td style="width:31.7pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{i + 1}</strong></span>
                            </p>
                        </td>
                        <td style="width:178.85pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{work.get("text", "")}</strong></span>
                            </p>
                        </td>
                        <td style="width:67.15pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{start_date} – {end_date} г.</strong></span>
                            </p>
                        </td>
                        <td style="width:152.2pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{work.get("additional_text", "")}</strong></span>
                            </p>
                        </td>
                        <td style="width:51.3pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{work.get("unit", "")}</strong></span>
                            </p>
                        </td>
                        <td style="width:51.3pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{work.get("quantity", "")}</strong></span>
                            </p>
                        </td>
                        <td style="width:65.45pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{work.get("sum", "")}</strong></span>
                            </p>
                        </td>
                        <td style="width:72.55pt; border:0.75pt solid #000000; padding:0pt 5.03pt; vertical-align:top">
                            <p style="line-height:115%; font-size:9pt">
                                <span style="font-family:Calibri"><strong>{work.get("amount", "")}</strong></span>
                            </p>
                        </td>
                    </tr>
                    '''
                    
                    # Суммируем для итогов
                    try:
                        # Суммируем количество
                        total_quantity += float(work.get("quantity", 0))
                        
                        # Суммируем цену за единицу
                        total_sum += float(work.get("sum", 0))
                        
                        # Суммируем общую стоимость
                        amount_str = work.get("amount", "0")
                        # Извлекаем число из строки с НДС
                        if "(с НДС:" in amount_str:
                            # Берем первое число (без НДС)
                            amount_str = amount_str.split("(с НДС:")[0].strip()
                        total_amount += float(amount_str)
                    except (ValueError, IndexError):
                        pass
                
                # Заменяем блок с работами
                # Находим и заменяем весь блок {% for work in works %} ... {% endfor %}
                import re
                pattern = r'{% for work in works %}(.*?){% endfor %}'
                html_content = re.sub(pattern, works_html, html_content, flags=re.DOTALL)
                
                # Итоговые значения
                # Показываем количество как целое число, если оно целое
                quantity_display = int(total_quantity) if total_quantity.is_integer() else total_quantity
                # Показываем сумму как целое число, если оно целое
                sum_display = int(total_sum) if total_sum.is_integer() else f'{total_sum:.2f}'
                html_content = html_content.replace('{{res_quantity}}', f'<strong>{quantity_display}</strong>')
                html_content = html_content.replace('{{res_sum}}', f'<strong>{sum_display}</strong>')
                html_content = html_content.replace('{{res}}', f'<strong>{total_amount:.2f}</strong>')
            preview_html = f'<div style="font-family: Times New Roman, serif; font-size: 14px; line-height: 1.5; color: #000;">{html_content}</div>'
            return JsonResponse({'success': True, 'preview_html': preview_html})
        else:
            # Для GET запроса показываем пустой шаблон
            html_content = html_content.replace('{{full_name}}', '')
            html_content = html_content.replace('{{phone_number}}', '')
            html_content = html_content.replace('{{iin}}', '')
            html_content = html_content.replace('{{doc_id}}', '')
            html_content = html_content.replace('{{act_id}}', '')
            html_content = html_content.replace('{{current_date}}', f'<strong>{timezone.now().strftime("%d.%m.%Y")}</strong>')
            html_content = html_content.replace('{{start_date}}', '')
            html_content = html_content.replace('{{end_date}}', '')
            html_content = html_content.replace('{{quantity}}', '')
            html_content = html_content.replace('{{sum}}', '')
            html_content = html_content.replace('{{amount}}', '')
            html_content = html_content.replace('{{text}}', '')
            html_content = html_content.replace('{{unit}}', '')
            html_content = html_content.replace('{{additional_text}}', '')
            html_content = html_content.replace('{{res_quantity}}', '')
            html_content = html_content.replace('{{res_sum}}', '')
            html_content = html_content.replace('{{res}}', '')
            preview_html = f'<div style="font-family: Times New Roman, serif; font-size: 14px; line-height: 1.5; color: #000;">{html_content}</div>'
            return JsonResponse({'success': True, 'preview_html': preview_html})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Ошибка загрузки документа: {str(e)}'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут сохранять документы
def gph_save(request):
    """Сохранение ГПХ договора и загрузка в Dropbox"""
    if request.method == 'POST':
        form = GPHContractForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            user = data['executor']
            start_date = data['start_date']
            end_date = data['end_date']
            # Согласующие
            approvers = []
            if 'approvers' in request.POST:
                approver_usernames = request.POST.getlist('approvers')
                User = get_user_model()
                for username in approver_usernames:
                    try:
                        appr = User.objects.get(username=username, role=3)
                        approvers.append({
                            'username': appr.username,
                            'full_name': appr.full_name,
                            'status': 'ожидание'
                        })
                    except User.DoesNotExist:
                        pass
            # Формируем actions
            actions = []
            actions.append({
                'role': 'executor',
                'username': user.username,
                'full_name': user.full_name,
                'status': 'ожидание'
            })
            actions.append({
                'role': 'initiator',
                'username': request.user.username,
                'full_name': request.user.get_full_name() or request.user.username,
                'status': 'ожидание'
            })
            for appr in approvers:
                actions.append({
                    'role': 'approver',
                    'username': appr['username'],
                    'full_name': appr['full_name'],
                    'status': 'ожидание'
                })
            signer = data['signer']
            actions.append({
                'role': 'signer',
                'username': signer.username,
                'full_name': signer.full_name,
                'status': 'ожидание'
            })
            # 1. Создаем запись в БД (генерируется doc_id)
            document = GphDocument.objects.create(
                doc_type='gph',
                user=user,
                full_name=user.full_name,
                start_date=start_date,
                end_date=end_date,
                file_path='',  # временно
                actions=actions
            )
            doc_id = document.doc_id
            # 2. Формируем docx с подстановкой всех данных
            from docx import Document as DocxDocument
            import io
            template_path = os.path.join(settings.BASE_DIR, 'static', 'docs', 'gph.docx')
            doc = DocxDocument(template_path)
            def replace_placeholders(text):
                text = text.replace('{{full_name}}', user.full_name)
                text = text.replace('{{doc_id}}', doc_id)
                text = text.replace('{{year}}', str(document.created_at.year))
                text = text.replace('{{current_date}}', document.created_at.strftime('%d.%m.%Y'))
                text = text.replace('{{start_date}}', start_date.strftime('%d.%m.%Y'))
                text = text.replace('{{end_date}}', end_date.strftime('%d.%m.%Y'))
                text = text.replace('{{nca_datas}}', 'Данные подписи НУЦ')
                text = text.replace('{{signer}}', signer.full_name)
                text = text.replace('{{signer_num}}', signer.proxy_number or '')
                text = text.replace('{{signer_date}}', signer.proxy_date.strftime('%d.%m.%Y') if signer.proxy_date else '')
                return text
            for paragraph in doc.paragraphs:
                paragraph.text = replace_placeholders(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.text = replace_placeholders(paragraph.text)
            # 3. Сохраняем docx в память
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            # 4. Загружаем в Dropbox
            access_token = get_dropbox_access_token()
            dbx = dropbox.Dropbox(access_token)
            dropbox_folder = f"/ncasign/{user.username}/"
            dropbox_filename = f"ГПХ-{user.username}-{document.created_at.strftime('%Y%m%d-%H%M%S')}.docx"
            dropbox_path = dropbox_folder + dropbox_filename
            dbx.files_upload(file_stream.read(), dropbox_path, mode=dropbox.files.WriteMode.overwrite)
            # 5. Получаем публичную ссылку
            shared_link_metadata = dbx.sharing_create_shared_link_with_settings(dropbox_path)
            public_url = shared_link_metadata.url.replace('?dl=0', '?raw=1')
            # 6. Сохраняем ссылку в модель
            document.file_path = public_url
            document.save()
            return JsonResponse({'success': True, 'file_path': public_url, 'doc_id': doc_id})
        else:
            return JsonResponse({'success': False, 'error': 'Неверные данные формы'})
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут скачивать акты
def act_download(request):
    """Скачивание акта в формате DOCX с подстановкой данных"""
    try:
        template_path = os.path.join(settings.BASE_DIR, 'static', 'docs', 'act.docx')
        if not os.path.exists(template_path):
            return JsonResponse({'success': False, 'error': 'Файл шаблона act.docx не найден'})
        
        # Создаем документ из шаблона
        doc = DocxDocument(template_path)
        
        # Получаем данные из запроса
        data = {
            'full_name': request.GET.get('full_name', ''),
            'phone_number': request.GET.get('phone_number', ''),
            'iin': request.GET.get('iin', ''),
            'doc_id': request.GET.get('doc_id', ''),
            'act_id': '',  # Оставляем пустым, будем генерировать на бэкенде
            'current_date': timezone.now().strftime('%d.%m.%Y'),
            'start_date': request.GET.get('start_date', ''),
            'end_date': request.GET.get('end_date', ''),
            'quantity': request.GET.get('quantity', ''),
            'unit_price': request.GET.get('unit_price', ''),
            'amount': request.GET.get('amount', ''),
            'created_date': '',  # Будем заполнять из ГПХ договора
        }
        
        # Если указан исполнитель, ищем его последний ГПХ договор только для ID
        executor_username = request.GET.get('executor', '')
        if executor_username and not data['doc_id']:
            try:
                User = get_user_model()
                user = User.objects.get(username=executor_username, role=4)
                # Ищем последний ГПХ договор пользователя только для получения ID
                last_gph = GphDocument.objects.filter(
                    user=user,
                    doc_type='gph'
                ).order_by('-created_at').first()
                
                if last_gph:
                    data['doc_id'] = last_gph.doc_id
                    data['created_date'] = last_gph.created_at.strftime('%d.%m.%Y') if last_gph.created_at else ''
            except User.DoesNotExist:
                pass
        
        # Вычисляем общую стоимость (количество × цена за единицу)
        try:
            quantity = float(data['quantity']) if data['quantity'] else 0
            unit_price = float(data['unit_price']) if data['unit_price'] else 0
            if quantity > 0 and unit_price > 0:
                data['amount'] = f"{quantity * unit_price:.2f}"
            else:
                data['amount'] = ''
        except (ValueError, ZeroDivisionError):
            data['amount'] = ''
        
        # Функция для замены плейсхолдеров
        def replace_placeholders(text):
            text = text.replace('{{full_name}}', data['full_name'])
            text = text.replace('{{phone_number}}', data['phone_number'])
            text = text.replace('{{iin}}', data['iin'])
            text = text.replace('{{doc_id}}', data['doc_id'])
            text = text.replace('{{created_date}}', data['created_date'])
            text = text.replace('{{act_id}}', data['act_id'])
            text = text.replace('{{current_date}}', data['current_date'])
            text = text.replace('{{start_date}}', data['start_date'])
            text = text.replace('{{end_date}}', data['end_date'])
            text = text.replace('{{quantity}}', data['quantity'])
            text = text.replace('{{sum}}', data['unit_price'])
            text = text.replace('{{amount}}', data['amount'])
            text = text.replace('{{res_quantitu}}', data['quantity'])
            text = text.replace('{{res}}', data['amount'])
            return text
        
        # Заменяем плейсхолдеры в параграфах
        for paragraph in doc.paragraphs:
            paragraph.text = replace_placeholders(paragraph.text)
        
        # Заменяем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = replace_placeholders(paragraph.text)
        
        # Сохраняем в память
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Создаем HTTP ответ для скачивания
        response = HttpResponse(
            file_stream.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="act.docx"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Ошибка создания документа: {str(e)}'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут получать данные пользователей
def user_data_api(request, username):
    """API для получения данных пользователя по username"""
    try:
        User = get_user_model()
        user = User.objects.get(username=username, role=4)  # Только исполнители
        
        user_data = {
            'full_name': user.full_name or '',
            'phone_number': user.phone_number or '',
            'iin': user.iin or '',
        }
        
        return JsonResponse({'success': True, 'user': user_data})
        
    except User.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Пользователь не найден'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Ошибка: {str(e)}'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут сохранять акты
def act_save(request):
    """Сохранение акта выполненных работ и загрузка в Dropbox"""
    if request.method == 'POST':
        form = ActForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            executor_username = data['executor']
            User = get_user_model()
            try:
                user = User.objects.get(username=executor_username, role=4)
            except User.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Исполнитель не найден'})
            last_gph = GphDocument.objects.filter(
                user=user, 
                doc_type='gph'
            ).order_by('-created_at').first()
            if not last_gph:
                return JsonResponse({'success': False, 'error': 'ГПХ договор не найден для данного исполнителя'})
            # Согласующие
            approvers = []
            if 'approvers' in request.POST:
                approver_usernames = request.POST.getlist('approvers')
                for username in approver_usernames:
                    try:
                        appr = User.objects.get(username=username, role=3)
                        approvers.append({
                            'username': appr.username,
                            'full_name': appr.full_name,
                            'status': 'ожидание'
                        })
                    except User.DoesNotExist:
                        pass
            # Формируем actions
            actions = []
            # 1. Исполнитель
            actions.append({
                'role': 'executor',
                'username': user.username,
                'full_name': user.full_name,
                'status': 'ожидание'
            })
            # 2. Инициатор
            actions.append({
                'role': 'initiator',
                'username': request.user.username,
                'full_name': request.user.get_full_name() or request.user.username,
                'status': 'ожидание'
            })
            # 3. Согласующие
            for appr in approvers:
                actions.append({
                    'role': 'approver',
                    'username': appr['username'],
                    'full_name': appr['full_name'],
                    'status': 'ожидание'
                })
            # 4. Подписант
            signer = data['signer']
            actions.append({
                'role': 'signer',
                'username': signer.username,
                'full_name': signer.full_name,
                'status': 'ожидание'
            })
            # Получаем массив работ из запроса
            import json
            works = []
            works_json = request.POST.get('works', '[]')
            try:
                works = json.loads(works_json)
            except json.JSONDecodeError:
                works = []
            
            # Если работ нет, создаем одну работу из данных формы
            if not works:
                try:
                    quantity = float(data['quantity']) if data['quantity'] else 0
                    unit_price = float(data['unit_price']) if data['unit_price'] else 0
                    amount = quantity * unit_price
                    if data.get('vat_included'):
                        amount_vat = amount * 1.12
                        amount_str = f"{amount:.2f} (с НДС: {amount_vat:.2f})"
                    else:
                        amount_str = f"{amount:.2f}"
                except (ValueError, ZeroDivisionError):
                    amount_str = "0.00"
                
                works = [{
                    'text': data['text'],
                    'start_date': data['start_date'].strftime('%Y-%m-%d') if data['start_date'] else '',
                    'end_date': data['end_date'].strftime('%Y-%m-%d') if data['end_date'] else '',
                    'additional_text': data['additional_text'],
                    'unit': data['unit'],
                    'quantity': data['quantity'],
                    'sum': data['unit_price'],
                    'amount': amount_str
                }]
            
            # Рассчитываем итоговые суммы
            total_quantity = 0
            total_sum = 0
            total_amount = 0
            
            for work in works:
                try:
                    total_quantity += float(work.get('quantity', 0))
                    total_sum += float(work.get('sum', 0))
                    amount_str = work.get('amount', '0')
                    if "(с НДС:" in amount_str:
                        amount_str = amount_str.split("(с НДС:")[0].strip()
                    total_amount += float(amount_str)
                except (ValueError, IndexError):
                    pass
            
            # Форматируем итоговые суммы
            total_quantity_str = str(int(total_quantity)) if total_quantity.is_integer() else str(total_quantity)
            total_sum_str = str(int(total_sum)) if total_sum.is_integer() else f"{total_sum:.2f}"
            total_amount_str = f"{total_amount:.2f}"
            
            document = ActDocument.objects.create(
                user=user,
                gph_document=last_gph,
                full_name=data['full_name'],
                phone_number=data['phone_number'],
                iin=data['iin'],
                start_date=data['start_date'],
                end_date=data['end_date'],
                works=works,
                total_quantity=total_quantity_str,
                total_sum=total_sum_str,
                total_amount=total_amount_str,
                file_path='',  # временно
                actions=actions
            )
            act_id = document.act_id
            # Загружаем файл в Dropbox
            try:
                template_path = os.path.join(settings.BASE_DIR, 'static', 'docs', 'act.docx')
                doc = DocxDocument(template_path)
                def replace_placeholders(text):
                    text = text.replace('{{full_name}}', data['full_name'])
                    text = text.replace('{{phone_number}}', data['phone_number'])
                    text = text.replace('{{iin}}', data['iin'])
                    text = text.replace('{{doc_id}}', last_gph.doc_id)
                    text = text.replace('{{created_date}}', last_gph.created_at.strftime('%d.%m.%Y') if last_gph.created_at else '')
                    text = text.replace('{{act_id}}', document.act_id)
                    text = text.replace('{{current_date}}', document.created_at.strftime('%d.%m.%Y'))
                    text = text.replace('{{start_date}}', data['start_date'].strftime('%d.%m.%Y') if hasattr(data['start_date'], 'strftime') else str(data['start_date']))
                    text = text.replace('{{end_date}}', data['end_date'].strftime('%d.%m.%Y') if hasattr(data['end_date'], 'strftime') else str(data['end_date']))
                    text = text.replace('{{res_quantity}}', total_quantity_str)
                    text = text.replace('{{res_sum}}', total_sum_str)
                    text = text.replace('{{res}}', total_amount_str)
                    return text
                
                # Заменяем плейсхолдеры в параграфах
                for paragraph in doc.paragraphs:
                    paragraph.text = replace_placeholders(paragraph.text)
                
                # Обрабатываем таблицы
                for table in doc.tables:
                    # Находим таблицу с работами (таблица с заголовками "Выполнено работ")
                    if len(table.rows) > 0:
                        # Проверяем, есть ли в первой строке заголовок "Выполнено работ"
                        first_row_text = ' '.join([cell.text for cell in table.rows[0].cells])
                        if 'Выполнено работ' in first_row_text:
                            # Находим строку с нумерацией столбцов (строка с цифрами 1, 2, 3, 4, 5, 6, 7, 8)
                            numbering_row = None
                            for i, row in enumerate(table.rows):
                                row_text = ' '.join([cell.text.strip() for cell in row.cells])
                                if '1' in row_text and '2' in row_text and '3' in row_text and '4' in row_text and '5' in row_text and '6' in row_text and '7' in row_text and '8' in row_text:
                                    numbering_row = i
                                    break
                            
                            if numbering_row is not None:
                                # Находим итоговую строку
                                total_row = None
                                for i, row in enumerate(table.rows):
                                    row_text = ' '.join([cell.text for cell in row.cells])
                                    if 'Итого' in row_text:
                                        total_row = i
                                        break
                                
                                if total_row is not None:
                                    # Удаляем все строки между нумерацией и итоговой строкой
                                    rows_to_remove = []
                                    for i in range(numbering_row + 1, total_row):
                                        rows_to_remove.append(i)
                                    
                                    # Удаляем строки в обратном порядке
                                    for i in reversed(rows_to_remove):
                                        table._element.remove(table.rows[i]._element)
                                    
                                    # Добавляем строки с работами перед итоговой строкой
                                    for i, work in enumerate(works):
                                        # Создаем новую строку
                                        new_row = table.add_row()
                                        
                                        # Заполняем данные
                                        if len(new_row.cells) >= 8:
                                            new_row.cells[0].text = str(i + 1)  # Номер
                                            new_row.cells[1].text = work.get('text', '')  # Наименование работ
                                            
                                            # Форматируем даты правильно
                                            start_date = work.get('start_date', '')
                                            end_date = work.get('end_date', '')
                                            if start_date and end_date:
                                                # Преобразуем даты в нужный формат
                                                try:
                                                    from datetime import datetime
                                                    if isinstance(start_date, str):
                                                        start_dt = datetime.strptime(start_date, '%Y-%m-%d')
                                                        start_formatted = start_dt.strftime('%d.%m.%Y')
                                                    else:
                                                        start_formatted = start_date.strftime('%d.%m.%Y') if hasattr(start_date, 'strftime') else str(start_date)
                                                    
                                                    if isinstance(end_date, str):
                                                        end_dt = datetime.strptime(end_date, '%Y-%m-%d')
                                                        end_formatted = end_dt.strftime('%d.%m.%Y')
                                                    else:
                                                        end_formatted = end_date.strftime('%d.%m.%Y') if hasattr(end_date, 'strftime') else str(end_date)
                                                    
                                                    new_row.cells[2].text = f"{start_formatted} – {end_formatted} г."
                                                except:
                                                    new_row.cells[2].text = f"{start_date} – {end_date} г."
                                            else:
                                                new_row.cells[2].text = ""
                                            
                                            new_row.cells[3].text = work.get('additional_text', '')  # Дополнительная информация
                                            new_row.cells[4].text = work.get('unit', 'месяц')  # Единица измерения
                                            new_row.cells[5].text = work.get('quantity', '')  # Количество
                                            new_row.cells[6].text = work.get('sum', '')  # Цена за единицу
                                            new_row.cells[7].text = work.get('amount', '')  # Стоимость
                    
                    # Заменяем плейсхолдеры в ячейках
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.text = replace_placeholders(paragraph.text)
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                access_token = get_dropbox_access_token()
                dbx = dropbox.Dropbox(access_token)
                dropbox_folder = f"/ncasign/{user.username}/"
                dropbox_filename = f"Акт-{user.username}-{document.created_at.strftime('%Y%m%d-%H%M%S')}.docx"
                dropbox_path = dropbox_folder + dropbox_filename
                dbx.files_upload(file_stream.read(), dropbox_path, mode=dropbox.files.WriteMode.overwrite)
                shared_link_metadata = dbx.sharing_create_shared_link_with_settings(dropbox_path)
                public_url = shared_link_metadata.url.replace('?dl=0', '?raw=1')
                document.file_path = public_url
                document.save()
            except Exception as e:
                document.delete()
                return JsonResponse({'success': False, 'error': f'Ошибка загрузки файла: {str(e)}'})
            # Добавляем акт в пакет (сессия)
            act_info = {
                'act_id': document.act_id,
                'executor_username': executor_username,
                'executor_name': data['full_name'],
                'amount': amount_str,
                'created_at': document.created_at.isoformat()
            }
            package_acts = request.session.get('package_acts', [])
            existing_executors = [act.get('executor_username') for act in package_acts]
            if executor_username in existing_executors:
                return JsonResponse({'success': False, 'error': 'Этот исполнитель уже добавлен в пакет'})
            package_acts.append(act_info)
            request.session['package_acts'] = package_acts
            request.session.modified = True
            return JsonResponse({
                'success': True, 
                'message': f'Акт для {data["full_name"]} добавлен в пакет',
                'package_count': len(package_acts),
                'act_id': document.act_id
            })
        else:
            return JsonResponse({'success': False, 'error': 'Неверные данные формы'})
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут добавлять акты в пакет
def add_act_to_package(request):
    """Добавление акта в пакет (сохранение акта + добавление в сессию)"""
    if request.method == 'POST':
        form = ActForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            executor_username = data['executor']
            User = get_user_model()
            try:
                user = User.objects.get(username=executor_username, role=4)
            except User.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Исполнитель не найден'})
            package_acts = request.session.get('package_acts', [])
            existing_executors = [act.get('executor_username') for act in package_acts]
            if executor_username in existing_executors:
                return JsonResponse({'success': False, 'error': 'Этот исполнитель уже добавлен в пакет'})
            last_gph = GphDocument.objects.filter(
                user=user, 
                doc_type='gph'
            ).order_by('-created_at').first()
            if not last_gph:
                return JsonResponse({'success': False, 'error': 'ГПХ договор не найден для данного исполнителя'})
            approvers = []
            if 'approvers' in request.POST:
                approver_usernames = request.POST.getlist('approvers')
                for username in approver_usernames:
                    try:
                        appr = User.objects.get(username=username, role=3)
                        approvers.append({
                            'username': appr.username,
                            'full_name': appr.full_name,
                            'status': 'ожидание'
                        })
                    except User.DoesNotExist:
                        pass
            # Определяем подписанта ДО формирования actions
            if 'signer' in data and data['signer']:
                signer = data['signer']
            else:
                signer = User.objects.filter(role=2).first()
                if not signer:
                    return JsonResponse({'success': False, 'error': 'Не найден подписант в системе'})
            # Теперь формируем actions
            actions = []
            actions.append({
                'role': 'executor',
                'username': user.username,
                'full_name': user.full_name,
                'status': 'ожидание'
            })
            actions.append({
                'role': 'initiator',
                'username': request.user.username,
                'full_name': request.user.get_full_name() or request.user.username,
                'status': 'ожидание'
            })
            for appr in approvers:
                actions.append({
                    'role': 'approver',
                    'username': appr['username'],
                    'full_name': appr['full_name'],
                    'status': 'ожидание'
                })
            actions.append({
                'role': 'signer',
                'username': signer.username,
                'full_name': signer.full_name,
                'status': 'ожидание'
            })
            # Получаем массив работ из запроса
            import json
            works = []
            works_json = request.POST.get('works', '[]')
            try:
                works = json.loads(works_json)
            except json.JSONDecodeError:
                works = []
            
            # Если работ нет, создаем одну работу из данных формы
            if not works:
                try:
                    quantity = float(data['quantity']) if data['quantity'] else 0
                    unit_price = float(data['unit_price']) if data['unit_price'] else 0
                    amount_calc = quantity * unit_price
                    if data.get('vat_included'):
                        amount_vat = amount_calc * 1.12
                        amount_str = f"{amount_calc:.2f} (с НДС: {amount_vat:.2f})"
                    else:
                        amount_str = f"{amount_calc:.2f}"
                except (ValueError, ZeroDivisionError):
                    amount_str = "0.00"
                
                works = [{
                    'text': data['text'],
                    'start_date': data['start_date'].strftime('%Y-%m-%d') if data['start_date'] else '',
                    'end_date': data['end_date'].strftime('%Y-%m-%d') if data['end_date'] else '',
                    'additional_text': data['additional_text'],
                    'unit': data['unit'],
                    'quantity': data['quantity'],
                    'sum': data['unit_price'],
                    'amount': amount_str
                }]
            
            # Рассчитываем итоговые суммы
            total_quantity = 0
            total_sum = 0
            total_amount = 0
            
            for work in works:
                try:
                    total_quantity += float(work.get('quantity', 0))
                    total_sum += float(work.get('sum', 0))
                    amount_str = work.get('amount', '0')
                    if "(с НДС:" in amount_str:
                        amount_str = amount_str.split("(с НДС:")[0].strip()
                    total_amount += float(amount_str)
                except (ValueError, IndexError):
                    pass
            
            # Форматируем итоговые суммы
            total_quantity_str = str(int(total_quantity)) if total_quantity.is_integer() else str(total_quantity)
            total_sum_str = str(int(total_sum)) if total_sum.is_integer() else f"{total_sum:.2f}"
            total_amount_str = f"{total_amount:.2f}"
            
            document = ActDocument.objects.create(
                user=user,
                gph_document=last_gph,
                full_name=data['full_name'],
                phone_number=data['phone_number'],
                iin=data['iin'],
                start_date=data['start_date'],
                end_date=data['end_date'],
                works=works,
                total_quantity=total_quantity_str,
                total_sum=total_sum_str,
                total_amount=total_amount_str,
                file_path='',  # временно
                actions=actions
            )
            
            # Загружаем файл в Dropbox
            try:
                template_path = os.path.join(settings.BASE_DIR, 'static', 'docs', 'act.docx')
                doc = DocxDocument(template_path)
                
                def replace_placeholders(text):
                    text = text.replace('{{full_name}}', data['full_name'])
                    text = text.replace('{{phone_number}}', data['phone_number'])
                    text = text.replace('{{iin}}', data['iin'])
                    text = text.replace('{{doc_id}}', last_gph.doc_id)
                    text = text.replace('{{created_date}}', last_gph.created_at.strftime('%d.%m.%Y') if last_gph.created_at else '')
                    text = text.replace('{{act_id}}', document.act_id)
                    text = text.replace('{{current_date}}', document.created_at.strftime('%d.%m.%Y'))
                    text = text.replace('{{start_date}}', data['start_date'].strftime('%d.%m.%Y') if hasattr(data['start_date'], 'strftime') else str(data['start_date']))
                    text = text.replace('{{end_date}}', data['end_date'].strftime('%d.%m.%Y') if hasattr(data['end_date'], 'strftime') else str(data['end_date']))
                    text = text.replace('{{res_quantity}}', total_quantity_str)
                    text = text.replace('{{res_sum}}', total_sum_str)
                    text = text.replace('{{res}}', total_amount_str)
                    return text
                
                # Заменяем плейсхолдеры в параграфах
                for paragraph in doc.paragraphs:
                    paragraph.text = replace_placeholders(paragraph.text)
                
                # Обрабатываем таблицы
                for table in doc.tables:
                    # Находим таблицу с работами (таблица с заголовками "Выполнено работ")
                    if len(table.rows) > 0:
                        # Проверяем, есть ли в первой строке заголовок "Выполнено работ"
                        first_row_text = ' '.join([cell.text for cell in table.rows[0].cells])
                        if 'Выполнено работ' in first_row_text:
                            # Находим строку с нумерацией столбцов (строка с цифрами 1, 2, 3, 4, 5, 6, 7, 8)
                            numbering_row = None
                            for i, row in enumerate(table.rows):
                                row_text = ' '.join([cell.text.strip() for cell in row.cells])
                                if '1' in row_text and '2' in row_text and '3' in row_text and '4' in row_text and '5' in row_text and '6' in row_text and '7' in row_text and '8' in row_text:
                                    numbering_row = i
                                    break
                            
                            if numbering_row is not None:
                                # Находим итоговую строку
                                total_row = None
                                for i, row in enumerate(table.rows):
                                    row_text = ' '.join([cell.text for cell in row.cells])
                                    if 'Итого' in row_text:
                                        total_row = i
                                        break
                                
                                if total_row is not None:
                                    # Удаляем все строки между нумерацией и итоговой строкой
                                    rows_to_remove = []
                                    for i in range(numbering_row + 1, total_row):
                                        rows_to_remove.append(i)
                                    
                                    # Удаляем строки в обратном порядке
                                    for i in reversed(rows_to_remove):
                                        table._element.remove(table.rows[i]._element)
                                    
                                    # Добавляем строки с работами перед итоговой строкой
                                    for i, work in enumerate(works):
                                        # Создаем новую строку
                                        new_row = table.add_row()
                                        
                                        # Заполняем данные
                                        if len(new_row.cells) >= 8:
                                            new_row.cells[0].text = str(i + 1)  # Номер
                                            new_row.cells[1].text = work.get('text', '')  # Наименование работ
                                            
                                            # Форматируем даты правильно
                                            start_date = work.get('start_date', '')
                                            end_date = work.get('end_date', '')
                                            if start_date and end_date:
                                                # Преобразуем даты в нужный формат
                                                try:
                                                    from datetime import datetime
                                                    if isinstance(start_date, str):
                                                        start_dt = datetime.strptime(start_date, '%Y-%m-%d')
                                                        start_formatted = start_dt.strftime('%d.%m.%Y')
                                                    else:
                                                        start_formatted = start_date.strftime('%d.%m.%Y') if hasattr(start_date, 'strftime') else str(start_date)
                                                    
                                                    if isinstance(end_date, str):
                                                        end_dt = datetime.strptime(end_date, '%Y-%m-%d')
                                                        end_formatted = end_dt.strftime('%d.%m.%Y')
                                                    else:
                                                        end_formatted = end_date.strftime('%d.%m.%Y') if hasattr(end_date, 'strftime') else str(end_date)
                                                    
                                                    new_row.cells[2].text = f"{start_formatted} – {end_formatted} г."
                                                except:
                                                    new_row.cells[2].text = f"{start_date} – {end_date} г."
                                            else:
                                                new_row.cells[2].text = ""
                                            
                                            new_row.cells[3].text = work.get('additional_text', '')  # Дополнительная информация
                                            new_row.cells[4].text = work.get('unit', 'месяц')  # Единица измерения
                                            new_row.cells[5].text = work.get('quantity', '')  # Количество
                                            new_row.cells[6].text = work.get('sum', '')  # Цена за единицу
                                            new_row.cells[7].text = work.get('amount', '')  # Стоимость
                    
                    # Заменяем плейсхолдеры в ячейках
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.text = replace_placeholders(paragraph.text)
                
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                
                access_token = get_dropbox_access_token()
                dbx = dropbox.Dropbox(access_token)
                dropbox_folder = f"/ncasign/{user.username}/"
                dropbox_filename = f"Акт-{user.username}-{document.created_at.strftime('%Y%m%d-%H%M%S')}.docx"
                dropbox_path = dropbox_folder + dropbox_filename
                dbx.files_upload(file_stream.read(), dropbox_path, mode=dropbox.files.WriteMode.overwrite)
                
                shared_link_metadata = dbx.sharing_create_shared_link_with_settings(dropbox_path)
                public_url = shared_link_metadata.url.replace('?dl=0', '?raw=1')
                
                document.file_path = public_url
                document.save()
                
            except Exception as e:
                # Если не удалось загрузить в Dropbox, удаляем запись
                document.delete()
                return JsonResponse({'success': False, 'error': f'Ошибка загрузки файла: {str(e)}'})
            
            # Добавляем акт в пакет (сессия)
            act_info = {
                'act_id': document.act_id,
                'executor_username': executor_username,
                'executor_name': data['full_name'],
                'amount': total_amount_str,
                'created_at': document.created_at.isoformat()
            }
            
            package_acts.append(act_info)
            request.session['package_acts'] = package_acts
            request.session.modified = True
            
            return JsonResponse({
                'success': True, 
                'message': f'Акт для {data["full_name"]} добавлен в пакет',
                'package_count': len(package_acts),
                'act_id': document.act_id
            })
        else:
            return JsonResponse({'success': False, 'error': 'Неверные данные формы'})
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут сохранять пакеты
def save_package(request):
    """Сохранение пакета актов в БД"""
    if request.method == 'POST':
        package_acts = request.session.get('package_acts', [])
        
        if len(package_acts) < 2:
            return JsonResponse({'success': False, 'error': 'Для создания пакета нужно минимум 2 акта'})
        
        try:
            # Создаем пакет
            package = ActPackage.objects.create(
                created_by=request.user,
                acts=[act['act_id'] for act in package_acts],
                status='ready'
            )
            
            # Очищаем сессию
            request.session.pop('package_acts', None)
            request.session.modified = True
            
            return JsonResponse({
                'success': True, 
                'message': f'Пакет {package.package_id} успешно создан с {len(package_acts)} актами',
                'package_id': package.package_id
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': f'Ошибка создания пакета: {str(e)}'})
    
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут очищать пакеты
def clear_package(request):
    """Очистка пакета из сессии"""
    if request.method == 'POST':
        package_acts = request.session.get('package_acts', [])
        
        if not package_acts:
            return JsonResponse({'success': False, 'error': 'Пакет уже пуст'})
        
        # Удаляем все акты из БД
        for act_info in package_acts:
            try:
                act = ActDocument.objects.get(act_id=act_info['act_id'])
                act.delete()
            except ActDocument.DoesNotExist:
                pass
        
        # Очищаем сессию
        request.session.pop('package_acts', None)
        request.session.modified = True
        
        return JsonResponse({
            'success': True, 
            'message': f'Пакет очищен, удалено {len(package_acts)} актов'
        })
    
    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})

@login_required
@role_required([1, 5])  # Только Админ и Редактор могут получать информацию о пакете
def get_package_info(request):
    """Получение информации о текущем пакете в сессии"""
    package_acts = request.session.get('package_acts', [])
    
    return JsonResponse({
        'success': True,
        'package_count': len(package_acts),
        'acts': package_acts
    })

@require_GET
@login_required
def act_package_count(request):
    """Возвращает количество актов в виртуальном пакете пользователя (из сессии)"""
    package_acts = request.session.get('package_acts', [])
    return JsonResponse({'success': True, 'count': len(package_acts)})
